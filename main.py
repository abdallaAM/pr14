import sympy as sp
from fpdf import FPDF
from docx import Document
import os

class DerivativeCalculator:
    def __init__(self):
        self.x, self.y, self.z = sp.symbols('x y z')

    def compute_derivative(self, expression, variable):
        expr = sp.sympify(expression)
        derivative = sp.diff(expr, variable)
        return sp.simplify(derivative)

    def save_to_pdf(self, expression, result):
        """Сохранение PDF с поиском шрифта."""
        pdf = FPDF()
        pdf.add_page()

        # Проверяем доступные шрифты
        fonts = {
            "dejavu": "DejaVuSans.ttf",
            "arial": "C:/Windows/Fonts/arial.ttf"
        }
        font_used = None

        for name, path in fonts.items():
            if os.path.exists(path):
                pdf.add_font(name, "", path, uni=True)
                font_used = name
                break

        if font_used is None:
            raise RuntimeError("Не найден подходящий шрифт!")

        pdf.set_font(font_used, size=12)
        pdf.multi_cell(190, 10, f"Выражение: {expression}\nПроизводная: {result}")
        pdf.output("result.pdf")

    def save_to_docx(self, expression, result):
        doc = Document()
        doc.add_paragraph(f"Выражение: {expression}")
        doc.add_paragraph(f"Производная: {result}")
        doc.save("result.docx")

# === ТЕСТЫ ===
if __name__ == "__main__":
    calc = DerivativeCalculator()

    expression = "x**3 + sin(x) * cos(x)"
    result = calc.compute_derivative(expression, calc.x)
    print(f"Производная {expression} по x: {result}")

    calc.save_to_pdf(expression, result)
    calc.save_to_docx(expression, result)

    print("✅ Файлы result.pdf и result.docx успешно созданы!")
